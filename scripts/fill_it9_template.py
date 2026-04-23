from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TEMPLATE_PATH = Path(r"C:\Users\kirlg\Downloads\IT9 Template\IT9aL_Documentation-Template.docx")
BACKUP_PATH = TEMPLATE_PATH.with_name(TEMPLATE_PATH.stem + "-backup.docx")

TITLE = "LocalLift PH: A Web-Based Marketplace Platform for Local Products and Verified Sellers"
SUBMITTED_BY = "[Replace with group members in ascending order by last name]"
MONTH_YEAR = "April 2026"

BACKGROUND_PARAGRAPHS = [
    (
        "Micro, small, and medium enterprises (MSMEs) remain a major part of the Philippine economy, "
        "but many local sellers still rely on fragmented selling channels such as personal social media "
        "pages, chat-based ordering, and manual transaction coordination. While these channels are easy "
        "to start with, they often create problems in product discovery, seller credibility, order "
        "tracking, and customer trust."
    ),
    (
        "The continuing expansion of digital connectivity in the Philippines shows that web-based systems "
        "can play a stronger role in improving market access. The Philippine Statistics Authority reported "
        "that 67.3 percent of individuals aged 10 years and over used the internet in 2024, indicating a "
        "large base of users who can participate in online services and digital commerce (Philippine "
        "Statistics Authority, 2025)."
    ),
    (
        "Government initiatives also show that digital transformation has become a strategic priority for "
        "local enterprises. The Department of Trade and Industry emphasized that MSMEs must adopt digital "
        "tools to improve competitiveness and expand market reach, and it reported onboarding 300,321 "
        "MSMEs nationwide to e-commerce channels from January 2023 to March 2024 (Department of Trade "
        "and Industry, 2024a). The DTI also continued MSME digitalization programs focused on digital "
        "payments and online selling capabilities in 2024 (Department of Trade and Industry, 2024b)."
    ),
    (
        "Despite this progress, many local sellers still do not have a dedicated platform where products, "
        "shop identity, customer communication, and transactions are managed in one organized system. "
        "Buyers often need to ask for details manually, compare products across scattered pages, and rely "
        "on informal messaging to place orders. This setup can cause inconsistent product information, "
        "slower transactions, and weak accountability."
    ),
    (
        "To address these conditions, the researchers propose LocalLift PH, a web-based marketplace "
        "platform for local products and verified sellers. The system provides buyer, seller, and "
        "administrator workflows. Buyers can browse products and shops, manage carts, place orders, "
        "maintain addresses, exchange messages with sellers, and submit reviews. Sellers can apply for "
        "approval, manage products, view orders and earnings, communicate with buyers, and update shop "
        "settings. Administrators can review seller applications and approve or reject product listings. "
        "Through these features, the system aims to create a more organized, transparent, and reliable "
        "online environment for local commerce."
    ),
]

GENERAL_PROBLEM = (
    "Local sellers and buyers often rely on fragmented and informal online selling processes that do not "
    "provide a centralized, trustworthy, and efficient marketplace for product discovery, seller "
    "verification, communication, and transaction management."
)

SPECIFIC_PROBLEMS = [
    "Buyers have difficulty finding local products in one organized platform because listings are often scattered across separate pages and informal channels.",
    "Buyers cannot easily assess seller legitimacy when there is no structured seller verification and product approval workflow.",
    "Manual or chat-based selling processes make inquiry, ordering, and follow-up slower and more prone to misunderstanding.",
    "Sellers lack an integrated environment for managing products, orders, earnings, customer messages, and shop information.",
    "Administrators need a more systematic way to review seller applications and moderate product listings before publication.",
    "The absence of built-in cart, checkout, address management, order history, and review features reduces the efficiency and credibility of the buying process.",
]

GENERAL_OBJECTIVE = (
    "To develop a web-based marketplace platform for local products that enables verified sellers to "
    "manage their stores and allows buyers to browse, communicate, and transact through a centralized and "
    "reliable system."
)

SPECIFIC_OBJECTIVES = [
    "To provide buyer-side features for account access, shop browsing, product viewing, cart management, checkout, address management, order tracking, and product reviews.",
    "To provide seller-side features for registration, application submission, store setup, product management, order monitoring, earnings viewing, messaging, and shop settings management.",
    "To implement an administrator workflow for reviewing seller applications and approving or rejecting product listings.",
    "To improve trust and marketplace quality by allowing only approved sellers and approved products to be visible to buyers.",
    "To support more organized buyer-seller communication through an integrated messaging module with text and image sharing.",
    "To create a more efficient and user-friendly digital environment that promotes local products and supports the online growth of small businesses.",
]

SCOPE_PARAGRAPHS = [
    (
        "The proposed system is a web-based marketplace platform focused on promoting local products through "
        "a centralized online environment. The system supports three primary user roles: buyers, sellers, "
        "and administrators."
    ),
    (
        "For buyers, the system includes account access, browsing of shops and products, viewing of product "
        "categories, adding items to a cart, selecting delivery addresses, placing orders, reviewing order "
        "history, cancelling eligible orders, starting conversations with sellers, sending chat messages, "
        "and submitting product reviews."
    ),
    (
        "For sellers, the system includes account registration, seller application submission, shop setup, "
        "product creation, product inventory and status management, viewing received orders, checking "
        "earnings summaries, managing buyer conversations, editing seller profile information, and "
        "maintaining shop settings."
    ),
    (
        "For administrators, the system includes login access, seller application review, seller approval "
        "status management, and product approval or rejection before listings become available to buyers."
    ),
    (
        "The system is implemented as a Laravel-based web application with a PHP backend, relational "
        "database support through Laravel migrations and models, and a frontend built using Blade templates, "
        "Tailwind CSS, and Vite-based asset management."
    ),
]

LIMITATIONS_PARAGRAPHS = [
    "The study is limited to a web application and does not include a dedicated mobile application for Android or iOS.",
    "The system does not currently include online payment gateway integration, automated courier tracking, advanced analytics, recommendation engines, or external logistics API integration.",
    "Seller verification and product approval depend on administrator review, so moderation quality is still influenced by human decision-making and processing time.",
    "The platform is intended for marketplace operations and does not cover broader enterprise functions such as accounting, taxation automation, warehouse optimization, or multi-branch business management.",
    "The effectiveness of the system also depends on internet access, proper user adoption, and the availability of updated product and seller information.",
]

FUNCTIONAL_REQUIREMENTS = [
    "The system shall allow buyers, sellers, and administrators to log in using their respective access flows.",
    "The system shall allow new seller accounts to register and submit application requirements for approval.",
    "The system shall allow administrators to approve or reject seller applications.",
    "The system shall allow sellers to add, edit, and manage product listings with details such as name, category, description, price, stock, dimensions, shipping fee, image, and status.",
    "The system shall allow administrators to approve or reject pending product listings.",
    "The system shall display only approved, active products from approved sellers to buyer-facing pages.",
    "The system shall allow buyers to browse products, categories, and seller shops.",
    "The system shall allow buyers to view product details and seller shop information.",
    "The system shall allow buyers to add products to a cart and update or remove cart items.",
    "The system shall prevent users from ordering their own products.",
    "The system shall allow buyers to select saved addresses and proceed to checkout.",
    "The system shall compute subtotal, shipping fee, and total order cost during checkout.",
    "The system shall allow buyers to place orders and store the order with itemized product details.",
    "The system shall allow buyers to view their order history and order details.",
    "The system shall allow buyers to cancel eligible orders and record cancellation reasons.",
    "The system shall allow sellers to view order information related to their listed products.",
    "The system shall allow sellers to view summary information such as sales, active products, pending orders, and conversation counts.",
    "The system shall allow buyers and sellers to start conversations and exchange messages within the platform.",
    "The system shall allow chat messages to include text and image attachments.",
    "The system shall allow users to manage profile information and seller shop settings.",
    "The system shall allow buyers to submit product reviews within the platform.",
]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def format_body(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = None


def format_list_item(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def replace_text(paragraphs, target: str, replacement: str) -> None:
    for paragraph in paragraphs:
        if paragraph.text.strip() == target:
            paragraph.text = replacement
            return


def ensure_text(paragraphs, target: str, replacement: str) -> None:
    for paragraph in paragraphs:
        if paragraph.text.strip() in {target, replacement}:
            paragraph.text = replacement
            return


def find_index(paragraphs, target: str, start: int = 0) -> int:
    for index in range(start, len(paragraphs)):
        if paragraphs[index].text.strip() == target:
            return index
    raise ValueError(f"Could not find paragraph: {target}")


def rebuild_statement_section(doc: Document, statement_idx: int, objectives_idx: int) -> None:
    paragraphs = doc.paragraphs
    for idx in range(objectives_idx - 1, statement_idx, -1):
        remove_paragraph(paragraphs[idx])

    anchor = doc.paragraphs[statement_idx]
    body_style = doc.paragraphs[statement_idx + 1].style if statement_idx + 1 < len(doc.paragraphs) else None

    p = insert_paragraph_after(anchor, GENERAL_PROBLEM, body_style)
    format_body(p)

    p = insert_paragraph_after(p, "Specific Problems:", body_style)
    format_body(p)
    for num, item in enumerate(SPECIFIC_PROBLEMS, start=1):
        p = insert_paragraph_after(p, f"{num}. {item}", body_style)
        format_list_item(p)


def rebuild_objectives_section(doc: Document, objectives_idx: int, scope_idx: int) -> None:
    paragraphs = doc.paragraphs
    for idx in range(scope_idx - 1, objectives_idx, -1):
        remove_paragraph(paragraphs[idx])

    anchor = doc.paragraphs[objectives_idx]
    body_style = doc.paragraphs[objectives_idx + 1].style if objectives_idx + 1 < len(doc.paragraphs) else None

    p = insert_paragraph_after(anchor, "General Objective:", body_style)
    format_body(p)
    p = insert_paragraph_after(p, GENERAL_OBJECTIVE, body_style)
    format_body(p)
    p = insert_paragraph_after(p, "Specific Objectives:", body_style)
    format_body(p)
    for num, item in enumerate(SPECIFIC_OBJECTIVES, start=1):
        p = insert_paragraph_after(p, f"{num}. {item}", body_style)
        format_list_item(p)


def rebuild_scope_section(doc: Document, scope_idx: int, functional_idx: int) -> None:
    paragraphs = doc.paragraphs
    for idx in range(functional_idx - 1, scope_idx, -1):
        remove_paragraph(paragraphs[idx])

    anchor = doc.paragraphs[scope_idx]
    body_style = doc.paragraphs[scope_idx + 1].style if scope_idx + 1 < len(doc.paragraphs) else None

    p = insert_paragraph_after(anchor, "Scope:", body_style)
    format_body(p)
    for item in SCOPE_PARAGRAPHS:
        p = insert_paragraph_after(p, item, body_style)
        format_body(p)

    p = insert_paragraph_after(p, "Limitations:", body_style)
    format_body(p)
    for item in LIMITATIONS_PARAGRAPHS:
        p = insert_paragraph_after(p, item, body_style)
        format_body(p)


def rebuild_functional_requirements(doc: Document, functional_idx: int) -> None:
    paragraphs = doc.paragraphs
    for idx in range(len(paragraphs) - 1, functional_idx, -1):
        remove_paragraph(paragraphs[idx])

    anchor = doc.paragraphs[functional_idx]
    body_style = doc.paragraphs[functional_idx + 1].style if functional_idx + 1 < len(doc.paragraphs) else None

    p = insert_paragraph_after(anchor, "The following functional requirements define the expected behavior of the proposed system:", body_style)
    format_body(p)
    for num, item in enumerate(FUNCTIONAL_REQUIREMENTS, start=1):
        p = insert_paragraph_after(p, f"{num}. {item}", body_style)
        format_list_item(p)


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    BACKUP_PATH.write_bytes(TEMPLATE_PATH.read_bytes())

    doc = Document(str(TEMPLATE_PATH))

    ensure_text(doc.paragraphs, "PROJECT TITLE", TITLE)
    ensure_text(doc.paragraphs, "Ascending order based on last name.", SUBMITTED_BY)
    ensure_text(doc.paragraphs, "MONTH YEAR", MONTH_YEAR)

    background_heading_idx = find_index(doc.paragraphs, "Background of Study", start=70)

    for idx in range(len(doc.paragraphs) - 1, background_heading_idx, -1):
        remove_paragraph(doc.paragraphs[idx])

    body_style = doc.paragraphs[background_heading_idx].style
    current = doc.paragraphs[background_heading_idx]

    for item in BACKGROUND_PARAGRAPHS:
        current = insert_paragraph_after(current, item, body_style)
        format_body(current)

    current = insert_paragraph_after(current, "Statement of the Problem", body_style)
    current = insert_paragraph_after(current, "General Problem:", body_style)
    format_body(current)
    current = insert_paragraph_after(current, GENERAL_PROBLEM, body_style)
    format_body(current)
    current = insert_paragraph_after(current, "Specific Problems:", body_style)
    format_body(current)
    for num, item in enumerate(SPECIFIC_PROBLEMS, start=1):
        current = insert_paragraph_after(current, f"{num}. {item}", body_style)
        format_list_item(current)

    current = insert_paragraph_after(current, "Objectives", body_style)
    current = insert_paragraph_after(current, "General Objective:", body_style)
    format_body(current)
    current = insert_paragraph_after(current, GENERAL_OBJECTIVE, body_style)
    format_body(current)
    current = insert_paragraph_after(current, "Specific Objectives:", body_style)
    format_body(current)
    for num, item in enumerate(SPECIFIC_OBJECTIVES, start=1):
        current = insert_paragraph_after(current, f"{num}. {item}", body_style)
        format_list_item(current)

    current = insert_paragraph_after(current, "Scope and Limitations", body_style)
    current = insert_paragraph_after(current, "Scope:", body_style)
    format_body(current)
    for item in SCOPE_PARAGRAPHS:
        current = insert_paragraph_after(current, item, body_style)
        format_body(current)
    current = insert_paragraph_after(current, "Limitations:", body_style)
    format_body(current)
    for item in LIMITATIONS_PARAGRAPHS:
        current = insert_paragraph_after(current, item, body_style)
        format_body(current)

    current = insert_paragraph_after(current, "Functional Requirements", body_style)
    current = insert_paragraph_after(current, "The following functional requirements define the expected behavior of the proposed system:", body_style)
    format_body(current)
    for num, item in enumerate(FUNCTIONAL_REQUIREMENTS, start=1):
        current = insert_paragraph_after(current, f"{num}. {item}", body_style)
        format_list_item(current)

    doc.save(str(TEMPLATE_PATH))


if __name__ == "__main__":
    main()
